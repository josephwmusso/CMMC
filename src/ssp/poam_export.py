"""
src/ssp/poam_export.py

Generate POA&M as PDF (fpdf2) and DOCX (python-docx).
"""

import os
import json
import datetime

from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import text, create_engine

DATABASE_URL = os.getenv(
    "DATABASE_URL",
    "postgresql://cmmc:localdev@localhost:5432/cmmc",
)

EXPORT_DIR = os.path.join("data", "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

RISK_ORDER = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2}


def _safe(text_val):
    """Latin-1 safe string for fpdf2 Helvetica."""
    if text_val is None:
        return ""
    s = str(text_val)
    s = s.replace("\u2014", "--").replace("\u2013", "-")
    s = s.replace("\u2018", "'").replace("\u2019", "'")
    s = s.replace("\u201c", '"').replace("\u201d", '"')
    s = s.replace("\u2026", "...").replace("\u2022", "*")
    s = s.replace("\u00b7", "*").replace("\r", "")
    return s.encode("latin-1", errors="replace").decode("latin-1")


def _fetch_poam_data(org_id):
    """Fetch POA&M items with control metadata."""
    engine = create_engine(DATABASE_URL)
    with engine.connect() as conn:
        org_row = conn.execute(
            text("SELECT name FROM organizations WHERE id = :oid"),
            {"oid": org_id},
        ).fetchone()
        org_name = org_row[0] if org_row else "Apex Defense Solutions"

        rows = conn.execute(text("""
            SELECT p.id, p.control_id, p.weakness_description, p.remediation_plan,
                   p.milestone_changes, p.risk_level, p.status::text,
                   p.scheduled_completion, c.title, c.family_abbrev, c.points
            FROM poam_items p
            JOIN controls c ON c.id = p.control_id
            WHERE p.org_id = :oid
            ORDER BY c.points DESC, p.scheduled_completion ASC, p.control_id
        """), {"oid": org_id}).fetchall()

    items = []
    counts = {"OPEN": 0, "IN_PROGRESS": 0, "CLOSED": 0, "OVERDUE": 0}
    pts_at_risk = 0

    for r in rows:
        status = r[6] or "OPEN"
        points = r[10] or 1
        counts[status] = counts.get(status, 0) + 1
        if status in ("OPEN", "IN_PROGRESS"):
            pts_at_risk += points
        items.append({
            "id": r[0], "control_id": r[1], "weakness": r[2] or "",
            "remediation": r[3] or "", "milestones": r[4],
            "risk": r[5] or "MEDIUM", "status": status,
            "deadline": r[7], "title": r[8] or "", "family": r[9] or "",
            "points": points,
        })

    # Sort: CRITICAL first, then HIGH, then MEDIUM; within same risk, by deadline
    items.sort(key=lambda x: (RISK_ORDER.get(x["risk"], 9), str(x["deadline"] or "9999")))

    return org_name, items, counts, pts_at_risk


# ═══════════════════════════════════════════════════════════════════════
# PDF Export
# ═══════════════════════════════════════════════════════════════════════

class POAMPDF(FPDF):
    def __init__(self, org_name):
        super().__init__()
        self.org_name = org_name
        self.set_auto_page_break(auto=True, margin=20)

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 8, f"POA&M - {self.org_name}", align="L")
        self.cell(0, 8, "CUI", align="R", new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(180, 180, 180)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(4)

    def footer(self):
        if self.page_no() == 1:
            return
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        now = datetime.datetime.now().strftime("%Y-%m-%d")
        self.cell(0, 10, f"POA&M - {self.org_name} - Page {self.page_no()}/{{nb}} - Generated {now}", align="C")


def generate_poam_pdf(org_id):
    """Generate POA&M PDF. Returns file path."""
    org_name, items, counts, pts_at_risk = _fetch_poam_data(org_id)

    pdf = POAMPDF(org_name)
    pdf.alias_nb_pages()
    pdf.set_margins(20, 20, 20)

    # Cover page
    pdf.add_page()
    pdf.ln(50)
    pdf.set_font("Helvetica", "B", 26)
    pdf.set_text_color(30, 30, 30)
    pdf.cell(0, 14, "Plan of Action & Milestones", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 18)
    pdf.cell(0, 10, "(POA&M)", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.set_font("Helvetica", "", 14)
    pdf.set_text_color(60, 60, 60)
    pdf.cell(0, 8, org_name, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, "NIST SP 800-171 Rev 2 / CMMC Level 2", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, f"{len(items)} Items | {pts_at_risk} Points at Risk", align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.set_y(-50)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(180, 40, 40)
    pdf.cell(0, 8, "CUI -- CONTROLLED UNCLASSIFIED INFORMATION", align="C", new_x="LMARGIN", new_y="NEXT")

    # Summary page
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(30, 60, 120)
    pdf.cell(0, 10, "POA&M Summary", new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(59, 130, 246)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
    pdf.ln(8)

    summary = [
        ("Total Items", str(len(items))),
        ("Open", str(counts.get("OPEN", 0))),
        ("In Progress", str(counts.get("IN_PROGRESS", 0))),
        ("Overdue", str(counts.get("OVERDUE", 0))),
        ("Closed", str(counts.get("CLOSED", 0))),
        ("Points at Risk", str(pts_at_risk)),
        ("Critical Items", str(sum(1 for i in items if i["risk"] == "CRITICAL"))),
        ("High Items", str(sum(1 for i in items if i["risk"] == "HIGH"))),
        ("Medium Items", str(sum(1 for i in items if i["risk"] == "MEDIUM"))),
    ]
    for label, value in summary:
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(80, 7, label)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 7, value, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # Items
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(30, 60, 120)
    pdf.cell(0, 10, "POA&M Items", new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(59, 130, 246)
    pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
    pdf.ln(6)

    for idx, item in enumerate(items, 1):
        if pdf.get_y() > pdf.h - 50:
            pdf.add_page()

        # Item header
        risk_colors = {"CRITICAL": (239, 68, 68), "HIGH": (245, 158, 11), "MEDIUM": (59, 130, 246)}
        rc = risk_colors.get(item["risk"], (100, 100, 100))

        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*rc)
        pdf.cell(12, 6, f"#{idx}")
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 6, f"{item['control_id']} -- {_safe(item['title'])}", new_x="LMARGIN", new_y="NEXT")

        # Status / Risk / Deadline
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*rc)
        pdf.cell(30, 5, item["risk"])
        pdf.set_text_color(100, 100, 100)
        pdf.cell(30, 5, item["status"])
        deadline_str = str(item["deadline"])[:10] if item["deadline"] else "Not set"
        pdf.cell(0, 5, f"Due: {deadline_str}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # Weakness
        if item["weakness"]:
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(60, 60, 60)
            pdf.cell(0, 5, "Weakness:", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(0, 4.5, _safe(item["weakness"]), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

        # Remediation
        if item["remediation"]:
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(60, 60, 60)
            pdf.cell(0, 5, "Remediation Plan:", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(0, 4.5, _safe(item["remediation"]), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

        # Separator
        pdf.set_draw_color(200, 200, 200)
        pdf.set_line_width(0.2)
        pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
        pdf.ln(4)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"POAM_{org_name.replace(' ', '_')}_{timestamp}.pdf"
    filepath = os.path.join(EXPORT_DIR, filename)
    pdf.output(filepath)
    return filepath


# ═══════════════════════════════════════════════════════════════════════
# DOCX Export
# ═══════════════════════════════════════════════════════════════════════

def generate_poam_docx(org_id):
    """Generate POA&M DOCX. Returns file path."""
    org_name, items, counts, pts_at_risk = _fetch_poam_data(org_id)

    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plan of Action & Milestones (POA&M)")
    run.bold = True
    run.font.size = Pt(22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(org_name).font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NIST SP 800-171 Rev 2 / CMMC Level 2")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONTROLLED UNCLASSIFIED INFORMATION")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(180, 40, 40)

    doc.add_page_break()

    # Summary
    doc.add_heading("POA&M Summary", level=1)
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Light Shading Accent 1'
    summary_data = [
        ("Total Items", str(len(items))),
        ("Open / In Progress", f"{counts.get('OPEN', 0)} / {counts.get('IN_PROGRESS', 0)}"),
        ("Overdue / Closed", f"{counts.get('OVERDUE', 0)} / {counts.get('CLOSED', 0)}"),
        ("Points at Risk", str(pts_at_risk)),
        ("Critical / High / Medium",
         f"{sum(1 for i in items if i['risk'] == 'CRITICAL')} / "
         f"{sum(1 for i in items if i['risk'] == 'HIGH')} / "
         f"{sum(1 for i in items if i['risk'] == 'MEDIUM')}"),
        ("Earliest Deadline", str(min((i["deadline"] for i in items if i["deadline"]), default="N/A"))[:10]),
    ]
    for idx, (label, value) in enumerate(summary_data):
        summary_table.cell(idx, 0).text = label
        summary_table.cell(idx, 1).text = value

    doc.add_paragraph()

    # POA&M Table
    doc.add_heading("POA&M Items", level=1)
    headers = ["#", "Control ID", "Title", "Weakness", "Remediation", "Risk", "Status", "Due Date"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)

    for idx, item in enumerate(items, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = item["control_id"]
        cells[2].text = item["title"]
        cells[3].text = (item["weakness"] or "")[:300]
        cells[4].text = (item["remediation"] or "")[:300]
        cells[5].text = item["risk"]
        cells[6].text = item["status"]
        cells[7].text = str(item["deadline"])[:10] if item["deadline"] else "N/A"

        # Size all cells
        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

        # Color risk cell
        risk_colors = {"CRITICAL": RGBColor(200, 50, 50), "HIGH": RGBColor(200, 120, 20), "MEDIUM": RGBColor(50, 100, 200)}
        rc = risk_colors.get(item["risk"])
        if rc:
            for run in cells[5].paragraphs[0].runs:
                run.font.color.rgb = rc
                run.bold = True

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("CUI -- Distribution authorized to U.S. Government agencies and their contractors.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"POAM_{org_name.replace(' ', '_')}_{timestamp}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)
    return filepath
